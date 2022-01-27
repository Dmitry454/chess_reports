from docx import Document
from datetime import datetime, date, time


name_of_counterpart = 'Павел Самсонов' #input()
my_name = 'Дмитрий Гантимуров'
I_played_white_first = input('Who played white in the first game(me/counterpart)?')
if I_played_white_first == 'me':
    Me_white_first = True
else:
    Me_white_first = False

# Creating a word document
document = Document()

# Creating a heading of the report
heading = document.add_heading('Отчет о результатах самостоятельной работы обучающегося по дисциплинам "Физическая культура" или "Элективные курсы по физической культуре и спорту"')
paragraph = document.add_paragraph(' ')
# Adding the first table
pers_info_table = document.add_table(rows=2, cols=5)
pers_info_table.style = 'LightShading-Accent2'

# filling the first table
information = ['Студ. Билет', 'ФИО', 'Группа', 'Спортивное отделение', 'Преподаватель',
'110503', 'Гантимуров Дмитрий Егорович', '1105', 'Шахматы', 'Иванов С.В.']
ind = 0
for row in pers_info_table.rows:
    for cell in row.cells:
        cell.text = information[ind]
        ind += 1

paragraph = document.add_paragraph(' ')

# Extracting data from file and filling arrays for each game(information_1 and information_2)
file = open("C:\\Users\\38812\\OneDrive\\Рабочий стол\\отчеты по шахматам\\ходы партии.txt", "r")
information_1 = []
information_2 = []
end_of_the_first_game = False
white_are_winners_in_first = False
white_are_winners_in_second = False

# каждая строка файла - отдельно взятый ход одной из сторон или номер хода
# проходим по файлу и записываем ходы в списки information_1 или information_2
for ind_of_line, line in enumerate(file):
    # если в строке присутствует символ '#' (индикатор того, что этим ходом ставится мат)
    if '#' in line:
        # по остатку от деления номера строки на 3 мы можем понять, кто поставил мат
        # если индекс при делении на 3 дает остаток 1, то мат поставили белые
        # следовательно, нам не нужно рассматривать следующую строку, поэтому
        # мы ставим флажок white_is_winners в положение False
        if ind_of_line % 3 == 1:
            if end_of_the_first_game == False:
                white_are_winners_in_first = True
            else:
                white_are_winners_in_second = True
        # Если флажок end_of_the_first_game равен False, то мы добавляем последний
        # ход белых в список information_1
        if end_of_the_first_game == False:
            information_1.append(line)
            end_of_the_first_game = True
        # иначе - в список information_2
        else:
            information_2.append(line)

    # Если в строке не содержится символа '#'
    else:
        if end_of_the_first_game == False:
            information_1.append(line)
        else:
            information_2.append(line)

#---------------------------------------
# THE FIRST game
#---------------------------------------
# Creating the table of names of players
table_of_names_of_the_first_game = document.add_table(rows=3, cols=1)

if Me_white_first == True:
    information = ['Шахматная партия', f'Белые: {my_name}', f'Черные: {name_of_counterpart}']
    Me_white_first = False
else:
    information = ['Шахматная партия', f'Белые: {name_of_counterpart}', f'Черные: {my_name}']
    Me_white_first = True

# Filling the table with information
ind = 0
for row in table_of_names_of_the_first_game.rows:
    for cell in row.cells:
        cell.text = information[ind]
        ind += 1

# Adding an empty paragraph
paragraph = document.add_paragraph(' ')

# Working out number of rows
if  len(information_1) % 3 == 0:
    number_of_rows_1 = len(information_1) // 3 + 1
else:
    number_of_rows_1 = len(information_1) // 3 + 2

# Creating the table of steps
first_game_table = document.add_table(rows=number_of_rows_1, cols=4)
first_game_table.style = 'LightShading-Accent1'

# Heading of the table of steps
cell_1 = first_game_table.cell(0, 0)
cell_1.text = '№'
cell_2 = first_game_table.cell(0, 1)
cell_2.text = 'Белые'
cell_3 = first_game_table.cell(0, 2)
cell_3.text = 'Черные'
cell_4 = first_game_table.cell(0, 3)
cell_4.text = 'Анализ ходов'

# Filling the body of the table of steps
for ind_of_row in range(number_of_rows_1 - 1):
    for ind_of_cell in range(0, 3):
        cell = first_game_table.cell(ind_of_row + 1, ind_of_cell)

        try:
            cell.text = information_1[ind_of_row*3 + ind_of_cell]

        except :
            # if white_are_winners_in_first == True:
            #         cell.text = 'Поражение'
            # else:
            #         cell.text = 'Победа'
            pass

# Adding an empty paragraph
paragraph = document.add_paragraph(' ')

#---------------------------------------
# SECOND game
#---------------------------------------

# Creating the table of names of players
table_of_names_of_the_second_game = document.add_table(rows=3, cols=1)

if Me_white_first == True:
    information = ['Шахматная партия', f'Белые: {my_name}', f'Черные: {name_of_counterpart}']
else:
    information = ['Шахматная партия', f'Белые: {name_of_counterpart}', f'Черные: {my_name}']

# Filling the table with information
ind = 0
for row in table_of_names_of_the_second_game.rows:
    for cell in row.cells:
        cell.text = information[ind]
        ind += 1

# Adding an empty paragraph
paragraph = document.add_paragraph(' ')

# Working out number of rows
if  len(information_2) % 2 == 0:
    number_of_rows_2 = len(information_2) // 3 + 1
else:
    number_of_rows_2 = len(information_2) // 3 + 2

# Creating the table of steps
second_game_table = document.add_table(rows=number_of_rows_2, cols=4)
second_game_table.style = 'LightShading-Accent1'

# Heading of the table of steps
cell_1 = second_game_table.cell(0, 0)
cell_1.text = '№'
cell_2 = second_game_table.cell(0, 1)
cell_2.text = 'Белые'
cell_3 = second_game_table.cell(0, 2)
cell_3.text = 'Черные'
cell_4 = second_game_table.cell(0, 3)
cell_4.text = 'Анализ ходов'

# Filling the body of the table of steps
for ind_of_row in range(number_of_rows_2 - 1):
    for ind_of_cell in range(0, 3):
        cell = second_game_table.cell(ind_of_row + 1, ind_of_cell)

        try:
            cell.text = information_2[ind_of_row*3 + ind_of_cell]

        except :
            pass
            # if ind_of_cell > 0 and ind_of_cell :
            #     cell.text = 'Поражение'
            # else:
            #     cell.text = 'Победа'

document.save(f'Otchet_TK_FViS_110503_{date.today().year}-{date.today().month}-{date.today().day}.docx')
